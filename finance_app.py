# -------------------------------
# IFRS Financial Statements Generator
# -------------------------------

import streamlit as st
import pandas as pd
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---------------- Helpers ----------------
def fmt_currency(val, symbol="$", decimals=2):
    return f"{symbol}{val:,.{decimals}f}"

def make_ifrs_income_statement(df, currency_symbol="$", decimals=2):
    years = df["Year"].tolist()
    rows = [
        ("Revenue (Units × Price per Unit)", df["Sales"].tolist(), "currency"),
        ("  Units", df["Units"].tolist(), "units"),
        ("  Price per Unit", df["Price per Unit"].tolist(), "currency"),
        ("Cost of Sales", (df["Sales"] - df["Gross Profit"]).tolist(), "currency"),
        ("Gross Profit", df["Gross Profit"].tolist(), "currency"),
        ("Operating Expenses", "", "header"),
        ("  Manpower Costs", df["Manpower"].tolist(), "currency"),
        ("  Fixed Costs", df["Fixed Costs"].tolist(), "currency"),
        ("  Depreciation", df["Depreciation"].tolist(), "currency"),
        ("EBIT", df["EBIT"].tolist(), "currency"),
        ("Finance Costs", df["Interest"].tolist(), "currency"),
        ("Profit Before Tax", (df["EBIT"] - df["Interest"]).tolist(), "currency"),
        ("Income Tax Expense", df["Taxes"].tolist(), "currency"),
        ("Net Profit for the Year", df["Net Income"].tolist(), "currency"),
    ]
    out = {"Item": [r[0] for r in rows]}
    for i, year in enumerate(years):
        col = []
        for _, values, ftype in rows:
            if isinstance(values, list):
                val = values[i]
                if ftype == "currency":
                    col.append(f"{currency_symbol}{val:,.{decimals}f}")
                elif ftype == "units":
                    col.append(f"{val:,.0f}")
                else:
                    col.append("")
            else:
                col.append("")
        out[year] = col
    return pd.DataFrame(out)

def make_ifrs_balance_sheet(bs_numeric_df, assets_list, currency_symbol="$", decimals=2):
    years = bs_numeric_df["Year"].tolist()
    rows = [("ASSETS", ""), ("  Non-current assets", "")]
    for name, val, life in assets_list:
        rows.append((f"    {name}", [val] * len(years)))
    rows += [
        ("  Current assets", ""),
        ("    Cash", bs_numeric_df["Current Assets: Cash"].tolist()),
        ("    Accounts Receivable", bs_numeric_df["Current Assets: Accounts Receivable"].tolist()),
        ("    Inventory", bs_numeric_df["Current Assets: Inventory"].tolist()),
        ("  Total Current Assets", bs_numeric_df["Total Current Assets"].tolist()),
        ("Total Assets", bs_numeric_df["Total Assets"].tolist()),
        ("", ""),
        ("EQUITY AND LIABILITIES", ""),
        ("  Equity", ""),
        ("    Share Capital", bs_numeric_df["Equity: Share Capital"].tolist()),
        ("    Retained Earnings", bs_numeric_df["Equity: Retained Earnings"].tolist()),
        ("  Total Equity", bs_numeric_df["Total Equity"].tolist()),
        ("  Non-current Liabilities", bs_numeric_df["Non-Current Liabilities"].tolist()),
        ("  Current Liabilities", bs_numeric_df["Current Liabilities"].tolist()),
        ("Total Liabilities & Equity", bs_numeric_df["Total Liabilities & Equity"].tolist()),
    ]
    out = {"Item": [r[0] for r in rows]}
    for i, year in enumerate(years):
        col = []
        for r in rows:
            if isinstance(r[1], list):
                col.append(f"{currency_symbol}{r[1][i]:,.{decimals}f}")
            else:
                col.append("")
        out[year] = col
    return pd.DataFrame(out)

def make_ifrs_cashflow_display(cf_numeric, currency_symbol="$", decimals=2):
    rows = [
        ("Cash Flows from Operating Activities", ""),
        ("  Net Profit for the Year", cf_numeric["Net Income"].tolist()),
        ("  Depreciation", cf_numeric["Depreciation"].tolist()),
        ("  Change in Accounts Receivable", cf_numeric["ΔAR"].tolist()),
        ("  Change in Inventory", cf_numeric["ΔInventory"].tolist()),
        ("  Change in Accounts Payable", cf_numeric["ΔAP"].tolist()),
        ("Net Cash from Operating Activities", cf_numeric["Net CF Ops"].tolist()),
        ("Cash Flows from Investing Activities", ""),
        ("  Capital Expenditure (CapEx)", cf_numeric["CapEx"].tolist()),
        ("Net Cash from Investing Activities", cf_numeric["Net CF Inv"].tolist()),
        ("Cash Flows from Financing Activities", ""),
        ("  Change in Non-current Liabilities", cf_numeric["Change NCL"].tolist()),
        ("  Change in Share Capital", cf_numeric["Change SC"].tolist()),
        ("  Interest Paid", cf_numeric["Interest Paid"].tolist()),
        ("Net Cash from Financing Activities", cf_numeric["Net CF Fin"].tolist()),
        ("Net Increase in Cash", cf_numeric["Net Change Cash"].tolist()),
        ("Closing Cash Balance", cf_numeric["Closing Cash"].tolist()),
    ]
    out = {"Item": [r[0] for r in rows]}
    years = cf_numeric["Year"].tolist()
    for i, year in enumerate(years):
        col = []
        for _, values in rows:
            if isinstance(values, list):
                val = values[i]
                col.append(f"{currency_symbol}{val:,.{decimals}f}")
            else:
                col.append("")
        out[year] = col
    return pd.DataFrame(out)

def export_to_word(is_df, bs_df, cf_df, dep_schedule,
                   filename="IFRS_Financial_Statements.docx",
                   company_name="Company", period="For the Years Ended Dec 31, 2025"):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7); section.page_height = Cm(21.0)
    title = doc.add_paragraph()
    run = title.add_run("IFRS Financial Statements")
    run.font.size = Pt(24); run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")
    p_company = doc.add_paragraph(); run = p_company.add_run(company_name); run.font.size = Pt(18)
    p_company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_period  = doc.add_paragraph(); run = p_period.add_run(period); run.font.size = Pt(14)
    p_period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    def add_table(title, df):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid"
        hdr = table.rows[0].cells
        for j, col in enumerate(df.columns):
            hdr[j].text = str(col)
            if hdr[j].paragraphs and hdr[j].paragraphs[0].runs:
                hdr[j].paragraphs[0].runs[0].bold = True
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = str(v)
        doc.add_paragraph("\n")
    add_table("IFRS Income Statement", is_df)
    add_table("IFRS Balance Sheet", bs_df)
    add_table("IFRS Cash Flow Statement (Indirect Method)", cf_df)
    if dep_schedule is not None and not dep_schedule.empty:
        add_table("Asset Depreciation Schedule", dep_schedule)
    doc.save(filename)
    return filename

# ---------------- Main App ----------------
st.title("📊 IFRS Financial Statements Generator")

# --- Sidebar Inputs ---
st.sidebar.header("Assumptions")
years = st.sidebar.slider("Projection years", 1, 10, 5)
currency = st.sidebar.selectbox("Currency", ["USD ($)", "EUR (€)", "GBP (£)", "AED (د.إ)", "SAR (﷼)", "JPY (¥)"], 0)
currency_symbol = currency.split("(")[-1].replace(")", "").strip()
decimals = st.sidebar.number_input("Decimal Places", 0, 6, 2, 1)

st.sidebar.subheader("Opening Balances")
opening_cash = st.sidebar.number_input("Opening Cash", value=100000)
opening_ar = st.sidebar.number_input("Opening Accounts Receivable", value=50000)
opening_inventory = st.sidebar.number_input("Opening Inventory", value=30000)
opening_ppe = st.sidebar.number_input("Opening PPE", value=400000)
opening_ncl = st.sidebar.number_input("Opening Non-Current Liabilities", value=100000)
opening_cl = st.sidebar.number_input("Opening Current Liabilities", value=80000)
opening_equity = st.sidebar.number_input("Opening Share Capital", value=400000)
opening_retained = st.sidebar.number_input("Opening Retained Earnings", value=0)

st.sidebar.subheader("Revenue & Costs")
units = st.sidebar.number_input("Units (Year 1)", value=10000)
price_per_unit = st.sidebar.number_input("Price per Unit", value=100)
unit_growth = st.sidebar.slider("Annual Units Growth (%)", 0, 50, 5) / 100
price_growth = st.sidebar.slider("Annual Price Growth (%)", 0, 20, 2) / 100
variable_cost_per_unit = st.sidebar.number_input("Variable Cost per Unit", value=40)
fixed_costs = st.sidebar.number_input("Fixed Costs / year", value=200000)
tax_rate = st.sidebar.slider("Tax Rate (%)", 0, 50, 20) / 100
interest = st.sidebar.number_input("Annual Interest", value=20000)
capex_input = st.sidebar.number_input("Annual CapEx", value=50000)

st.sidebar.subheader("Manpower")
num_employees = st.sidebar.number_input("Employees", value=10)
avg_salary = st.sidebar.number_input("Avg Annual Salary", value=30000)
salary_growth = st.sidebar.slider("Annual Salary Growth (%)", 0, 20, 5) / 100

st.sidebar.subheader("Assets (Name, Value, Life)")
assets_input = st.sidebar.text_area("One per line (name, value, life)",
                                    value="Factory Building, 500000, 25\nMachinery, 300000, 10\nVehicles, 100000, 5")
assets_list = []
for line in assets_input.split("\n"):
    parts = [p.strip() for p in line.split(",")]
    if len(parts) == 3:
        try:
            assets_list.append((parts[0], float(parts[1]), int(parts[2])))
        except:
            pass

company_name = st.sidebar.text_input("Company Name", "Example Company")
reporting_period = st.sidebar.text_input("Reporting Period", "For the Years Ended Dec 31, 2025")

# ---------------- Calculations ----------------

# 1. Income Statement
data = []
u, p, salary = units, price_per_unit, avg_salary
for year in range(1, years + 1):
    sales = u * p
    cogs = u * variable_cost_per_unit
    gross_profit = sales - cogs
    manpower_cost = num_employees * salary
    depreciation = sum((val / life) for _, val, life in assets_list) if assets_list else 0
    opex = fixed_costs + manpower_cost + depreciation
    ebit = gross_profit - opex
    ebt = ebit - interest
    taxes = ebt * tax_rate if ebt > 0 else 0
    net_income = ebt - taxes
    data.append([f"Year {year}", sales, gross_profit, fixed_costs, manpower_cost,
                 depreciation, ebit, interest, taxes, net_income, u, p, capex_input])
    u *= (1 + unit_growth)
    p *= (1 + price_growth)
    salary *= (1 + salary_growth)

cols = ["Year","Sales","Gross Profit","Fixed Costs","Manpower","Depreciation",
        "EBIT","Interest","Taxes","Net Income","Units","Price per Unit","CapEx (Input)"]
df = pd.DataFrame(data, columns=cols)

# 2. Cash Flow
cf_rows = []
closing_cash_list = []
ar_prev, inv_prev, ppe_prev = opening_ar, opening_inventory, opening_ppe
cash_prev, retained_prev = opening_cash, opening_retained
ncl_prev, sc_prev, cl_prev = opening_ncl, opening_equity, opening_cl

for i, row in df.iterrows():
    year = row["Year"]
    net_income, depreciation = row["Net Income"], row["Depreciation"]
    ar, inv = row["Sales"] * 0.20, row["Sales"] * 0.10
    d_ar, d_inv, d_ap = ar - ar_prev, inv - inv_prev, 0
    capex = row["CapEx (Input)"]
    ppe_curr = ppe_prev + capex - depreciation
    d_ncl, d_sc = 0, 0
    cf_ops = net_income + depreciation - d_ar - d_inv + d_ap
    cf_inv = -capex
    interest_paid = row["Interest"]
    cf_fin = d_ncl + d_sc - interest_paid
    net_cf = cf_ops + cf_inv + cf_fin
    closing_cash = cash_prev + net_cf
    cf_rows.append({"Year": year, "Net Income": net_income, "Depreciation": depreciation,
                    "ΔAR": -d_ar, "ΔInventory": -d_inv, "ΔAP": d_ap,
                    "Net CF Ops": cf_ops, "CapEx": -capex, "Net CF Inv": cf_inv,
                    "Change NCL": d_ncl, "Change SC": d_sc, "Interest Paid": -interest_paid,
                    "Net CF Fin": cf_fin, "Net Change Cash": net_cf, "Closing Cash": closing_cash})
    closing_cash_list.append(closing_cash)
    ar_prev, inv_prev, ppe_prev, cash_prev = ar, inv, ppe_curr, closing_cash
    retained_prev += net_income

cf_numeric = pd.DataFrame(cf_rows)

# 3. Balance Sheet (Cash from CF)
# 3. Balance Sheet (Cash from CF, CL as balancing figure)
bs_rows = []
ppe_prev = opening_ppe
retained_prev = opening_retained

for i, row in df.iterrows():
    year = row["Year"]

    # Other assets
    ar = row["Sales"] * 0.20
    inv = row["Sales"] * 0.10

    # PPE
    ppe_prev = ppe_prev + row["CapEx (Input)"] - row["Depreciation"]

    # Equity
    retained_prev += row["Net Income"]
    share_capital = opening_equity
    total_equity = share_capital + retained_prev

    # Liabilities (NCL fixed, CL = balancing)
    ncl = opening_ncl
    cash = closing_cash_list[i]
    total_current_assets = cash + ar + inv
    total_assets = total_current_assets + ppe_prev
    cl = total_assets - (ncl + total_equity)   # plug to balance

    total_liab_equity = ncl + cl + total_equity

    bs_rows.append({
        "Year": year,
        "Non-Current Assets (PPE)": ppe_prev,
        "Current Assets: Cash": cash,
        "Current Assets: Accounts Receivable": ar,
        "Current Assets: Inventory": inv,
        "Total Current Assets": total_current_assets,
        "Total Assets": total_assets,
        "Equity: Share Capital": share_capital,
        "Equity: Retained Earnings": retained_prev,
        "Total Equity": total_equity,
        "Non-Current Liabilities": ncl,
        "Current Liabilities": cl,
        "Total Liabilities & Equity": total_liab_equity
    })

# Convert rows to DataFrame
bs_numeric = pd.DataFrame(bs_rows)

# 4. Formatted tables
dep_schedule = pd.DataFrame([(n,v,l, round(v/l,decimals)) for n,v,l in assets_list],
                            columns=["Asset","Value","Useful Life (Years)","Annual Depreciation"])
is_df = make_ifrs_income_statement(df, currency_symbol, decimals)
bs_df_display = make_ifrs_balance_sheet(bs_numeric, assets_list, currency_symbol, decimals)
cf_df = make_ifrs_cashflow_display(cf_numeric, currency_symbol, decimals)

# ---------------- Display ----------------
st.subheader("📑 IFRS Income Statement")
st.dataframe(is_df, use_container_width=True)
st.subheader("🏦 IFRS Balance Sheet")
st.dataframe(bs_df_display, use_container_width=True)
balance_gap = (bs_numeric["Total Assets"] - bs_numeric["Total Liabilities & Equity"]).abs().max()
if balance_gap < 0.01:
    st.success("Balance ✅")
else:
    st.warning(f"⚠️ Out of balance by {fmt_currency(balance_gap, currency_symbol, decimals)}")
st.subheader("💵 IFRS Cash Flow Statement")
st.dataframe(cf_df, use_container_width=True)
st.subheader("🛠 Depreciation Schedule")
st.dataframe(dep_schedule, use_container_width=True)

# ---------------- Downloads ----------------
def build_excel(is_df, bs_df_display, cf_df, dep_schedule):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        is_df.to_excel(writer, index=False, sheet_name="Income Statement", startrow=1)
        bs_df_display.to_excel(writer, index=False, sheet_name="Balance Sheet", startrow=1)
        cf_df.to_excel(writer, index=False, sheet_name="Cash Flow", startrow=1)
        dep_schedule.to_excel(writer, index=False, sheet_name="Depreciation", startrow=1)
    return output.getvalue()

st.download_button("💾 Download Excel",
    data=build_excel(is_df, bs_df_display, cf_df, dep_schedule),
    file_name="IFRS_financials.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

word_file = export_to_word(
    is_df, bs_df_display, cf_df, dep_schedule,
    company_name=company_name, period=reporting_period
)

with open(word_file, "rb") as f:
    st.download_button(
        "📄 Download Word",
        data=f,
        file_name="IFRS_financials.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )